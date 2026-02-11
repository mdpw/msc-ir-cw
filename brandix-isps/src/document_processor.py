from docx import Document
import re
from typing import List, Dict

class DocumentProcessor:
    def __init__(self):
        self.strategic_objectives = []
        self.action_items = []
    
    def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""
    
    def load_strategic_plan(self, file_path: str):
        """Load and parse strategic plan"""
        try:
            doc = Document(file_path)
            self.strategic_objectives = self._parse_strategic_objectives(doc)
            print(f"Extracted {len(self.strategic_objectives)} strategic objectives")
        except Exception as e:
            print(f"Error loading strategic plan: {e}")
            self.strategic_objectives = []
        return self.strategic_objectives
    
    def load_action_plan(self, file_path: str):
        """Load and parse action plan"""
        try:
            doc = Document(file_path)
            self.action_items = self._parse_action_items(doc)
            print(f"Extracted {len(self.action_items)} action items")
        except Exception as e:
            print(f"Error loading action plan: {e}")
            self.action_items = []
        return self.action_items
    
    def _parse_strategic_objectives(self, doc) -> List[Dict]:
        """Parse strategic objectives from the Strategic Plan document"""
        objectives = []
        current_pillar = "Unknown"
        obj_id = 1
        
        # Convert all paragraphs to list for easier processing
        all_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        in_objectives_section = False
        i = 0
        
        while i < len(all_text):
            text = all_text[i]
            
            # Start collecting when we see "5. STRATEGIC PILLARS AND OBJECTIVES"
            if re.match(r'^5\.?\s*STRATEGIC PILLARS', text, re.IGNORECASE):
                in_objectives_section = True
                i += 1
                continue
            
            # Stop at next major section
            if in_objectives_section and re.match(r'^[6-9]\.?\s*[A-Z]', text):
                break
            
            if not in_objectives_section:
                i += 1
                continue
            
            # Detect PILLAR headings
            if text.startswith("PILLAR") and ":" in text:
                current_pillar = text.split(":", 1)[1].strip()
                i += 1
                continue
            
            # Pattern 1: Numbered sub-sections (1.1, 1.2, 2.1, etc.)
            if re.match(r'^\d+\.\d+\s+[A-Z]', text):
                objectives.append({
                    'id': f"OBJ-{obj_id:03d}",
                    'pillar': current_pillar,
                    'text': text,
                    'type': 'strategic_sub_section'
                })
                obj_id += 1
            
            # Pattern 2: Goal statements
            elif text.startswith("Goal:"):
                goal_text = text.replace("Goal:", "").strip()
                if len(goal_text) > 10:
                    objectives.append({
                        'id': f"OBJ-{obj_id:03d}",
                        'pillar': current_pillar,
                        'text': goal_text,
                        'type': 'goal'
                    })
                    obj_id += 1
            
            # Pattern 3: Key Initiatives (numbered list after "Key Initiatives:")
            elif "Key Initiatives:" in text:
                # Look ahead for numbered items
                j = i + 1
                while j < len(all_text) and re.match(r'^\d+\.\s+', all_text[j]):
                    initiative = re.sub(r'^\d+\.\s+', '', all_text[j])
                    if len(initiative) > 10:
                        objectives.append({
                            'id': f"OBJ-{obj_id:03d}",
                            'pillar': current_pillar,
                            'text': initiative,
                            'type': 'initiative'
                        })
                        obj_id += 1
                    j += 1
                i = j - 1  # Skip the processed items
            
            # Pattern 4: KPIs (bullet points after "KPIs:")
            elif text == "KPIs:":
                j = i + 1
                while j < len(all_text):
                    kpi_text = all_text[j]
                    # Check for bullet points or dashes
                    if kpi_text.startswith(("-", "•")) or ": " in kpi_text:
                        kpi_clean = kpi_text.lstrip("-•").strip()
                        if len(kpi_clean) > 10 and "by 20" in kpi_clean:  # KPIs usually have targets
                            objectives.append({
                                'id': f"OBJ-{obj_id:03d}",
                                'pillar': current_pillar,
                                'text': kpi_clean,
                                'type': 'kpi'
                            })
                            obj_id += 1
                        j += 1
                    else:
                        break
                i = j - 1
            
            # Pattern 5: Vision/Mission statements in pillar sections
            elif text.startswith("Vision:") and in_objectives_section:
                vision_text = text.replace("Vision:", "").strip()
                if len(vision_text) > 15:
                    objectives.append({
                        'id': f"OBJ-{obj_id:03d}",
                        'pillar': current_pillar,
                        'text': vision_text,
                        'type': 'vision'
                    })
                    obj_id += 1
            
            i += 1
        
        return objectives
    
    def _parse_action_items(self, doc) -> List[Dict]:
        """Parse action items from the Action Plan document"""
        actions = []
        current_pillar = "Unknown"
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect pillar sections
            pillar_match = re.match(r'PILLAR\s+\d+:\s*(.+)', text)
            if pillar_match:
                current_pillar = pillar_match.group(1).strip()
                continue
            
            # Detect ACTION items (ACTION ENV-AIR-001:, etc.)
            action_match = re.match(r'ACTION\s+([\w-]+):\s*(.+)', text)
            if action_match:
                action_id = action_match.group(1)
                action_title = action_match.group(2).strip()
                
                actions.append({
                    'id': action_id,
                    'title': action_title,
                    'pillar': current_pillar,
                    'text': action_title,
                    'type': 'action'
                })
        
        return actions
    
    def get_summary(self) -> Dict:
        """Get summary statistics"""
        return {
            'total_objectives': len(self.strategic_objectives),
            'total_actions': len(self.action_items),
            'objectives_by_type': self._count_by_type(self.strategic_objectives),
            'actions_by_pillar': self._count_by_pillar(self.action_items)
        }
    
    def _count_by_type(self, items: List[Dict]) -> Dict:
        """Count items by type"""
        counts = {}
        for item in items:
            item_type = item.get('type', 'unknown')
            counts[item_type] = counts.get(item_type, 0) + 1
        return counts
    
    def _count_by_pillar(self, items: List[Dict]) -> Dict:
        """Count items by pillar"""
        counts = {}
        for item in items:
            pillar = item.get('pillar', 'Unknown')
            counts[pillar] = counts.get(pillar, 0) + 1
        return counts
    
    def save_parsed_data(self, output_file: str):
        """Save parsed data for inspection"""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("BRANDIX DOCUMENT PROCESSING RESULTS\n")
            f.write("="*80 + "\n\n")
            
            summary = self.get_summary()
            f.write(f"Total Strategic Objectives: {summary['total_objectives']}\n")
            f.write(f"Total Action Items: {summary['total_actions']}\n\n")
            
            if summary['objectives_by_type']:
                f.write("Objectives by Type:\n")
                for obj_type, count in summary['objectives_by_type'].items():
                    f.write(f"  {obj_type}: {count}\n")
                f.write("\n")
            
            if summary['actions_by_pillar']:
                f.write("Actions by Pillar:\n")
                for pillar, count in summary['actions_by_pillar'].items():
                    f.write(f"  {pillar}: {count}\n")
                f.write("\n")
            
            if self.strategic_objectives:
                f.write("="*80 + "\n")
                f.write("STRATEGIC OBJECTIVES (All)\n")
                f.write("="*80 + "\n\n")
                for i, obj in enumerate(self.strategic_objectives, 1):
                    f.write(f"{i}. [{obj['id']}] {obj['pillar']}\n")
                    f.write(f"   Type: {obj['type']}\n")
                    f.write(f"   Text: {obj['text']}\n\n")
            
            if self.action_items:
                f.write("="*80 + "\n")
                f.write("ACTION ITEMS (All)\n")
                f.write("="*80 + "\n\n")
                for i, action in enumerate(self.action_items, 1):
                    f.write(f"{i}. [{action['id']}] {action['pillar']}\n")
                    f.write(f"   Title: {action['title']}\n\n")
        
        print(f"Parsed data saved to {output_file}")

# Test
if __name__ == "__main__":
    import os
    
    # Create outputs directory if it doesn't exist
    os.makedirs('outputs', exist_ok=True)
    
    processor = DocumentProcessor()
    
    print("="*80)
    print("BRANDIX DOCUMENT PROCESSOR TEST")
    print("="*80)
    
    print("\nLoading Strategic Plan...")
    objectives = processor.load_strategic_plan('data/BRANDIX_STRATEGIC_PLAN_2025.docx')
    
    print("\nLoading Action Plan...")
    actions = processor.load_action_plan('data/BRANDIX_ACTION_PLAN_YEAR_1.docx')
    
    print("\n" + "="*80)
    summary = processor.get_summary()
    print("SUMMARY:")
    print(f"  Total Objectives: {summary['total_objectives']}")
    print(f"  Total Actions: {summary['total_actions']}")
    
    if summary['objectives_by_type']:
        print("\n  Objectives by Type:")
        for obj_type, count in summary['objectives_by_type'].items():
            print(f"    {obj_type}: {count}")
    
    print("="*80)
    
    # Save detailed results
    processor.save_parsed_data('outputs/parsed_data.txt')
    print("\nProcessing complete! Check outputs/parsed_data.txt for details")
    
    # Show samples
    if objectives:
        print("\nSample Strategic Objectives (first 5):")
        for i, obj in enumerate(objectives[:5], 1):
            print(f"\n{i}. [{obj['id']}] {obj['type']}")
            print(f"   Pillar: {obj['pillar']}")
            print(f"   Text: {obj['text'][:80]}...")
    
    if actions:
        print("\n\nSample Action Items (first 5):")
        for i, action in enumerate(actions[:5], 1):
            print(f"\n{i}. [{action['id']}]")
            print(f"   Pillar: {action['pillar']}")
            print(f"   Title: {action['title']}")